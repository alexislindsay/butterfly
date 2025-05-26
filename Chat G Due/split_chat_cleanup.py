import fitz  # PyMuPDF for PDF handling
from docx import Document
from docx.shared import Inches

# Load the PDF
pdf_path = r"C:\Users\alexi\Downloads\35 Day Manifestation Journal.pdf"
pdf_document = fitz.open(pdf_path)

# Create a new Word document
doc = Document()

# Set the desired margins for the Word document (in inches)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Loop through each page in the PDF and insert it as an image into the Word document
for page_number in range(len(pdf_document)):
    # Extract the page as an image
    page = pdf_document.load_page(page_number)
    pix = page.get_pixmap()
    image_path = f"page_{page_number+1}.png"
    pix.save(image_path)

    # Add a new paragraph to create some space
    doc.add_paragraph()

    # Insert the image into the Word document, resizing to fit within the margins
    doc.add_picture(image_path, width=Inches(7.25))

    # Add a page break after each page except the last one
    if page_number < len(pdf_document) :
        doc.add_page_break()

# Save the Word document
doc.save("output_document.docx")

print("Document created successfully.")
